import os
import sys
from pathlib import Path
from datetime import datetime
import mimetypes
import hashlib
import json

def get_file_hash(filepath, algorithm='sha256'):
    """Calculate file hash"""
    hash_func = hashlib.new(algorithm)
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b''):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def get_basic_metadata(filepath):
    """Extract basic file system metadata"""
    stat = os.stat(filepath)
    path = Path(filepath)
    
    metadata = {
        'filename': path.name,
        'full_path': str(path.absolute()),
        'extension': path.suffix,
        'size_bytes': stat.st_size,
        'size_kb': round(stat.st_size / 1024, 2),
        'size_mb': round(stat.st_size / (1024 * 1024), 2),
        'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
        'mime_type': mimetypes.guess_type(filepath)[0],
        'is_symlink': path.is_symlink(),
        'sha256': get_file_hash(filepath, 'sha256'),
        'md5': get_file_hash(filepath, 'md5'),
    }
    
    if hasattr(stat, 'st_birthtime'):
        metadata['birth_time'] = datetime.fromtimestamp(stat.st_birthtime).isoformat()
    
    return metadata

def get_image_metadata(filepath):
    """Extract image metadata using Pillow and piexif"""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS, GPSTAGS
        
        metadata = {}
        with Image.open(filepath) as img:
            metadata['format'] = img.format
            metadata['mode'] = img.mode
            metadata['size'] = img.size
            metadata['width'] = img.width
            metadata['height'] = img.height
            
            if hasattr(img, 'info'):
                metadata['info'] = img.info
            
            # Extract EXIF data
            exif_data = img.getexif()
            if exif_data:
                exif = {}
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        value = value.decode(errors='ignore')
                    exif[tag] = value
                
                # Get GPS info if available
                if 'GPSInfo' in exif:
                    gps_info = {}
                    for key in exif['GPSInfo'].keys():
                        decode = GPSTAGS.get(key, key)
                        gps_info[decode] = exif['GPSInfo'][key]
                    exif['GPSInfo'] = gps_info
                
                metadata['exif'] = exif
        
        return metadata
    except ImportError:
        return {'error': 'PIL (Pillow) not installed. Install with: pip install Pillow'}
    except Exception as e:
        return {'error': str(e)}

def get_audio_metadata(filepath):
    """Extract audio metadata using mutagen"""
    try:
        from mutagen import File
        
        audio = File(filepath)
        if audio is None:
            return {'error': 'Unsupported audio format'}
        
        metadata = {
            'length_seconds': audio.info.length if hasattr(audio, 'info') else None,
            'bitrate': audio.info.bitrate if hasattr(audio, 'info') else None,
            'sample_rate': audio.info.sample_rate if hasattr(audio, 'info') else None,
            'channels': audio.info.channels if hasattr(audio, 'info') else None,
            'tags': dict(audio.tags) if audio.tags else {}
        }
        
        return metadata
    except ImportError:
        return {'error': 'mutagen not installed. Install with: pip install mutagen'}
    except Exception as e:
        return {'error': str(e)}

def get_video_metadata(filepath):
    """Extract video metadata using opencv"""
    try:
        import cv2
        
        cap = cv2.VideoCapture(filepath)
        metadata = {
            'frame_count': int(cap.get(cv2.CAP_PROP_FRAME_COUNT)),
            'fps': cap.get(cv2.CAP_PROP_FPS),
            'width': int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
            'height': int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
            'codec': int(cap.get(cv2.CAP_PROP_FOURCC)),
            'duration_seconds': int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) / cap.get(cv2.CAP_PROP_FPS) if cap.get(cv2.CAP_PROP_FPS) > 0 else 0
        }
        cap.release()
        
        return metadata
    except ImportError:
        return {'error': 'opencv-python not installed. Install with: pip install opencv-python'}
    except Exception as e:
        return {'error': str(e)}

def get_pdf_metadata(filepath):
    """Extract PDF metadata using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(filepath)
        metadata = {
            'num_pages': len(reader.pages),
            'metadata': dict(reader.metadata) if reader.metadata else {},
            'is_encrypted': reader.is_encrypted
        }
        
        # Extract text from first page as sample
        if len(reader.pages) > 0:
            metadata['first_page_text_sample'] = reader.pages[0].extract_text()[:500]
        
        return metadata
    except ImportError:
        return {'error': 'PyPDF2 not installed. Install with: pip install PyPDF2'}
    except Exception as e:
        return {'error': str(e)}

def get_docx_metadata(filepath):
    """Extract DOCX metadata using python-docx"""
    try:
        from docx import Document
        
        doc = Document(filepath)
        core_props = doc.core_properties
        
        metadata = {
            'author': core_props.author,
            'category': core_props.category,
            'comments': core_props.comments,
            'content_status': core_props.content_status,
            'created': core_props.created.isoformat() if core_props.created else None,
            'identifier': core_props.identifier,
            'keywords': core_props.keywords,
            'language': core_props.language,
            'last_modified_by': core_props.last_modified_by,
            'last_printed': core_props.last_printed.isoformat() if core_props.last_printed else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'revision': core_props.revision,
            'subject': core_props.subject,
            'title': core_props.title,
            'version': core_props.version,
            'num_paragraphs': len(doc.paragraphs),
            'num_tables': len(doc.tables),
        }
        
        return metadata
    except ImportError:
        return {'error': 'python-docx not installed. Install with: pip install python-docx'}
    except Exception as e:
        return {'error': str(e)}

def extract_all_metadata(filepath):
    """Extract all possible metadata from a file"""
    if not os.path.exists(filepath):
        return {'error': f'File not found: {filepath}'}
    
    # Get basic metadata
    metadata = {'basic': get_basic_metadata(filepath)}
    
    # Determine file type and extract specific metadata
    mime_type = metadata['basic']['mime_type']
    extension = metadata['basic']['extension'].lower()
    
    if mime_type and mime_type.startswith('image/'):
        metadata['image'] = get_image_metadata(filepath)
    
    if mime_type and mime_type.startswith('audio/'):
        metadata['audio'] = get_audio_metadata(filepath)
    
    if mime_type and mime_type.startswith('video/'):
        metadata['video'] = get_video_metadata(filepath)
    
    if extension == '.pdf' or mime_type == 'application/pdf':
        metadata['pdf'] = get_pdf_metadata(filepath)
    
    if extension == '.docx':
        metadata['docx'] = get_docx_metadata(filepath)
    
    return metadata

def main():
    if len(sys.argv) < 2:
        print("Usage: python script.py <filepath>")
        print("\nOptional libraries for extended metadata:")
        print("  pip install Pillow          # For image metadata")
        print("  pip install mutagen         # For audio metadata")
        print("  pip install opencv-python   # For video metadata")
        print("  pip install PyPDF2          # For PDF metadata")
        print("  pip install python-docx     # For DOCX metadata")
        sys.exit(1)
    
    filepath = sys.argv[1]
    metadata = extract_all_metadata(filepath)
    
    # Pretty print JSON
    print(json.dumps(metadata, indent=2, default=str))

if __name__ == '__main__':
    main()